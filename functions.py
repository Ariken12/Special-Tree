
import docx
import random as rand

def save_in_word(dots):
    document = docx.Document()
    document.add_heading('Построение невидимой елки :")', 0)
    document.add_paragraph('На деревянной(можно взять любой материал, который выдержит вес шаров) доске просверлить отверстия по заданным координатам и подписать их номерами из таблицы. После необходимо протянуть через заданное отверстие леску заданной длинны.')
    table = document.add_table(rows=len(dots)+1, cols=5)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Номер отверстия'
    table.cell(0, 1).text = 'Координаты по иксу'
    table.cell(0, 2).text = 'Координаты по игреку'
    table.cell(0, 3).text = 'Длинна лески'
    table.cell(0, 4).text = 'Расстояние до центра доски'
    summ = 0
    for i in range(len(dots)):
        table.cell(i+1, 0).text = str(i+1) 
        table.cell(i+1, 1).text = str(dots[i][0]) + ' мм'
        table.cell(i+1, 2).text = str(dots[i][1]) + ' мм'
        table.cell(i+1, 3).text = str(round(dots[i][2])) + ' мм'
        table.cell(i+1, 4).text = str(round(dots[i][3], 2)) + ' мм'
        summ += round(dots[i][2])
    document.add_paragraph('Всего понадобиться {x} мм ленты'.format(x=summ))
    document.save('Координаты отверстий и длинна лески.docx')


def random_color():
    global rand
    a = str(hex(int(rand.randint(0, 16777215)))).replace('0x', '')
    while len(a) < 6:
        a = '0' + a
    a = '0' * (6 - len(a)) + a 
    return '#' + a